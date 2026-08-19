import copy
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def format_run(run, size=12, bold=False, italic=False, color="000000", font="Times New Roman"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_center_paragraph(doc, text="", size=12, bold=False, italic=False, color="000000", space_after=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if text:
        run = paragraph.add_run(text)
        format_run(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_spacer(doc, points):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.add_run("")
    return paragraph


def set_cell_border(cell, color="000000", size="6", val="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def clear_cell_borders(cell):
    set_cell_border(cell, color="FFFFFF", size="0", val="nil")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_top_ornament(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(2.0), Cm(12.0), Cm(2.0)]
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width
        clear_cell_borders(cell)
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for index, symbol in [(0, "❦"), (2, "❦")]:
        paragraph = table.cell(0, index).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(symbol)
        format_run(run, size=32, bold=False, color="3F3F3F", font="Segoe UI Symbol")

    center = table.cell(0, 1)
    p1 = center.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size in [
        ("ỦY BAN NHÂN DÂN TỈNH ĐẮK LẮK", 15),
        ("\nBTC CUỘC THI SÁNG TẠO DÀNH CHO", 14),
        ("\nTHANH THIẾU NIÊN, NHI ĐỒNG TỈNH ĐẮK LẮK", 14),
        ("\nLẦN THỨ I NĂM 2026", 13),
    ]:
        run = p1.add_run(text)
        format_run(run, size=size, bold=True)
    p2 = center.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("──────  ❖  ──────")
    format_run(run, size=11, bold=True, color="444444")


def add_info_block(doc):
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(10.5)
        for cell in row.cells:
            clear_cell_borders(cell)
            set_cell_margins(cell, top=35, bottom=35, start=60, end=60)

    rows = [
        ("Tác giả/nhóm tác giả:", "Ksơr Nay Quang"),
        ("Sinh năm:", "....................        Điện thoại: ...................."),
        ("Địa chỉ:", "Trường PTDTNT THPT Đam San, xã Ea Drông, tỉnh Đắk Lắk"),
        ("Người hướng dẫn:", "Ngô Minh Tuấn"),
        ("Đơn vị thực hiện:", "Trường PTDTNT THPT Đam San"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        p_label = row.cells[0].paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_label.add_run(label)
        format_run(run, size=12.5, bold=True)
        p_value = row.cells[1].paragraphs[0]
        run = p_value.add_run(value)
        format_run(run, size=12.5, bold=True if "...." not in value else False)


def add_bottom_ribbon(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.2), Cm(8.8), Cm(3.2)]
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width
        clear_cell_borders(cell)
        set_cell_margins(cell, top=35, bottom=35, start=40, end=40)

    for index, symbol in [(0, "❧"), (2, "❦")]:
        p = table.cell(0, index).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(symbol)
        format_run(run, size=28, color="3F3F3F", font="Segoe UI Symbol")

    center = table.cell(0, 1)
    set_cell_border(center, color="666666", size="8", val="single")
    set_cell_shading(center, "E6E6E6")
    p = center.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Đắk Lắk, tháng 06/2026")
    format_run(run, size=12.5, bold=True, italic=True, color="222222")


def add_first_page_border(doc):
    section = doc.sections[0]
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is not None:
        sect_pr.remove(pg_borders)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    pg_borders.set(qn("w:display"), "firstPage")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "double")
        node.set(qn("w:sz"), "12")
        node.set(qn("w:space"), "20")
        node.set(qn("w:color"), "333333")
        pg_borders.append(node)
    sect_pr.append(pg_borders)


def build_cover_doc():
    cover = Document()
    section = cover.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    add_top_ornament(cover)
    add_spacer(cover, 72)

    add_center_paragraph(cover, "BÁO CÁO THUYẾT MINH", size=19, bold=True, space_after=2)
    add_center_paragraph(cover, "GIẢI PHÁP DỰ THI", size=18, bold=True, space_after=28)

    add_center_paragraph(cover, "Tên giải pháp:", size=22, bold=True, space_after=12)
    add_center_paragraph(cover, "DAM SAN GREEN", size=28, bold=True, color="146B4A", space_after=4)
    add_center_paragraph(
        cover,
        "Ứng dụng số hỗ trợ bảo vệ môi trường học đường nội trú",
        size=14,
        italic=True,
        color="555555",
        space_after=70,
    )

    add_info_block(cover)
    add_spacer(cover, 80)
    add_bottom_ribbon(cover)
    cover.add_page_break()
    return cover


def main():
    path = Path(sys.argv[1])
    backup = path.with_name(path.stem + ".backup_truoc_tao_bia.docx")
    if not backup.exists():
        shutil.copy2(path, backup)

    doc = Document(str(path))
    first_text = doc.paragraphs[0].text.strip() if doc.paragraphs else ""
    if "ỦY BAN NHÂN DÂN TỈNH ĐẮK LẮK" in first_text:
        print("Cover already exists; no changes made.")
        return

    cover = build_cover_doc()
    body = doc._body._element
    first = body[0]
    for element in list(cover._body._element):
        if element.tag == qn("w:sectPr"):
            continue
        first.addprevious(copy.deepcopy(element))

    add_first_page_border(doc)
    doc.save(str(path))
    print(f"saved={path}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
