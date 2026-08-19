from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\DamSanGreen")
OUTPUT = ROOT / "output" / "documents" / "Huong_dan_su_dung_va_tai_khoan_DamSanGreen.docx"
LOGO = ROOT / "app" / "src" / "main" / "res" / "drawable" / "logo_damsan_green.png"

GREEN = "0B7A4B"
MINT = "DFF7EC"
MINT_2 = "F1FCF7"
CYAN = "00A6B8"
INK = "1B263B"
MUTED = "596579"
GOLD = "F4C542"
RED = "D9534F"
BORDER = "BFE6D5"
LIGHT = "F5F7FA"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK, size: float = 10.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = BORDER, size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=90, start=120, bottom=90, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_body(doc: Document, text: str, after: int = 6, color: str = INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_note_box(doc: Document, title: str, body: str, fill: str = MINT_2, accent: str = GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=BORDER, size="10")
    set_cell_margins(table, top=130, bottom=130, start=170, end=170)
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(accent)
    p2 = table.cell(0, 0).add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_step_table(doc: Document, title: str, steps: list[tuple[str, str]]) -> None:
    add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", size="0")
    set_cell_margins(table, top=120, bottom=120, start=120, end=120)
    for idx, (step_title, detail) in enumerate(steps):
        cell = table.cell(0, idx)
        set_cell_shading(cell, [MINT, "E7F7FB", "FFF7D6", "F1FCF7"][idx % 4])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(str(idx + 1))
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor.from_string(GREEN)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(step_title)
        r2.bold = True
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = RGBColor.from_string(INK)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(detail)
        r3.font.size = Pt(9.2)
        r3.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_two_col_table(doc: Document, title: str, rows: list[tuple[str, str]], widths=(1.75, 4.75)) -> None:
    add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_margins(table)
    set_table_width(table, list(widths))
    set_cell_shading(table.cell(0, 0), GREEN)
    set_cell_shading(table.cell(0, 1), GREEN)
    set_cell_text(table.cell(0, 0), "Mục", bold=True, color="FFFFFF")
    set_cell_text(table.cell(0, 1), "Nội dung / thao tác", bold=True, color="FFFFFF")
    for label, value in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], label, bold=True, color=GREEN, size=10)
        set_cell_text(row.cells[1], value, color=INK, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(5)


def add_account_table(doc: Document, grade: int) -> None:
    add_heading(doc, f"Tài khoản học sinh khối {grade}", 3)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_margins(table, top=70, bottom=70, start=100, end=100)
    set_table_width(table, [0.78, 2.25, 1.65, 1.82])
    headers = ["Lớp", "Email đăng nhập", "Mật khẩu", "Ghi chú"]
    for idx, text in enumerate(headers):
        set_cell_shading(table.cell(0, idx), GREEN)
        set_cell_text(table.cell(0, idx), text, bold=True, color="FFFFFF", size=9.5)
    for i in range(1, 7):
        cls = f"{grade}A{i}"
        row = table.add_row()
        values = [cls, f"{cls.lower()}@damsan.edu.vn", "........................", "Tài khoản lớp"]
        for idx, value in enumerate(values):
            set_cell_text(row.cells[idx], value, bold=(idx == 0), color=INK, size=9.2)
            if i % 2 == 0:
                set_cell_shading(row.cells[idx], "FAFEFC")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, GREEN, 14, 8),
        ("Heading 2", 13, CYAN, 10, 5),
        ("Heading 3", 11.5, GREEN, 7, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc: Document) -> None:
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(4)
        p_logo.add_run().add_picture(str(LOGO), width=Cm(3.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("HƯỚNG DẪN SỬ DỤNG\nVÀ TÀI KHOẢN ĐĂNG NHẬP")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("Dam San Green - Ứng dụng bảo vệ môi trường học đường nội trú")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor.from_string(MUTED)

    add_note_box(
        doc,
        "Dùng để nộp kèm sản phẩm",
        "Tài liệu này hướng dẫn cài đặt APK, đăng nhập tài khoản lớp, thao tác gửi báo cáo rác, vào khu vực admin và quản lý bản tin/branding. Mật khẩu trong bảng cần được điền theo mật khẩu thật đã tạo trong Firebase Authentication.",
        fill=MINT_2,
    )

    meta = [
        ("Tên ứng dụng", "Dam San Green"),
        ("Đơn vị", "Trường PTDTNT THPT Đam San"),
        ("Địa điểm", "Xã Ea Drông, tỉnh Đắk Lắk"),
        ("Phiên bản hướng dẫn", "1.0 - cập nhật 29/06/2026"),
        ("File APK", "app-debug.apk hoặc DamSanGreen-v1.0.apk trong USB nộp kèm"),
    ]
    add_two_col_table(doc, "Thông tin nhanh", meta)
    doc.add_page_break()


def build_doc() -> None:
    doc = Document()
    setup_styles(doc)

    add_cover(doc)

    add_heading(doc, "1. Yêu cầu trước khi sử dụng", 1)
    add_two_col_table(
        doc,
        "Thiết bị và kết nối",
        [
            ("Thiết bị", "Điện thoại Android hoặc máy ảo Android trong Android Studio."),
            ("Internet", "Cần mạng để đăng nhập Firebase, upload ảnh Cloudinary, tải bản tin và dùng AI."),
            ("Quyền truy cập", "Cho phép Camera để chụp rác; cho phép Vị trí/GPS để lưu tọa độ báo cáo."),
            ("Tài khoản", "Tài khoản lớp dùng email dạng 11a1@damsan.edu.vn; tài khoản admin dùng email riêng đã gán role admin trong Firebase."),
        ],
    )
    add_note_box(
        doc,
        "Lưu ý về mật khẩu",
        "Firebase Authentication không cho ứng dụng đọc lại mật khẩu đã tạo. Vì vậy bảng tài khoản ở cuối tài liệu để sẵn ô mật khẩu: hãy điền đúng mật khẩu thật trong Firebase trước khi in hoặc nộp cho ban giám khảo.",
        fill="FFF8E1",
        accent="9A6B00",
    )

    add_heading(doc, "2. Cài đặt ứng dụng", 1)
    add_step_table(
        doc,
        "Luồng cài APK",
        [
            ("Mở file APK", "Chọn file trong USB/thư mục nộp"),
            ("Cho phép cài đặt", "Bật cài từ nguồn này nếu Android hỏi"),
            ("Mở ứng dụng", "Tìm biểu tượng Dam San Green"),
            ("Cấp quyền", "Camera, GPS khi app yêu cầu"),
        ],
    )
    add_body(doc, "Nếu chạy trên máy ảo Android Studio, kéo thả APK vào emulator hoặc dùng lệnh adb install với file APK trong thư mục nộp.")

    doc.add_page_break()
    add_heading(doc, "3. Đăng nhập tài khoản lớp", 1)
    add_step_table(
        doc,
        "Luồng đăng nhập học sinh",
        [
            ("Chọn lớp", "Ví dụ 11A1"),
            ("Nhập mật khẩu", "Mật khẩu lớp do giáo viên/admin cấp"),
            ("Bấm Đăng nhập", "App tự chuyển vào trang chính"),
            ("Kiểm tra lớp", "Hồ sơ và báo cáo phải hiện đúng lớp"),
        ],
    )
    add_two_col_table(
        doc,
        "Quy tắc tài khoản lớp",
        [
            ("Email", "App tự tạo email từ lớp được chọn: 11A1 -> 11a1@damsan.edu.vn."),
            ("Mật khẩu", "Nhập mật khẩu tương ứng với tài khoản đó trong Firebase Authentication."),
            ("Quên mật khẩu", "Bấm Quên mật khẩu nếu email reset đã được cấu hình trong Firebase."),
            ("Lỗi lớp Unknown", "Dùng bản APK mới; app đã tự suy ra lớp từ email và tự vá profile nếu thiếu className."),
        ],
    )

    add_heading(doc, "4. Gửi báo cáo rác", 1)
    add_step_table(
        doc,
        "Luồng học sinh báo cáo",
        [
            ("Quét rác", "Bấm QUÉT RÁC ở trang chính"),
            ("Nhập tên", "Ghi tên người báo cáo/minh chứng"),
            ("Chụp ảnh", "Chụp rõ rác hoặc khu vực vừa dọn"),
            ("Gửi báo cáo", "AI phân tích, upload ảnh và lưu điểm"),
        ],
    )
    add_two_col_table(
        doc,
        "Chi tiết thao tác báo cáo",
        [
            ("Ảnh minh chứng", "Nên chụp rõ chai nhựa, túi nilon, giấy, lon hoặc khu vực đã vệ sinh."),
            ("GPS", "Nếu có quyền vị trí, app lưu tọa độ báo cáo để hiển thị trên bản đồ."),
            ("AI duyệt rác", "AI ước lượng loại rác/khối lượng; báo cáo đủ tin cậy có thể tự duyệt, còn lại chờ admin."),
            ("Điểm", "Mặc định +10 điểm; báo cáo có khối lượng lớn có thể được cộng cao hơn theo AI."),
        ],
    )

    add_heading(doc, "5. Chế độ khách và bản đồ", 1)
    add_two_col_table(
        doc,
        "Khách tham quan",
        [
            ("Vào chế độ khách", "Tại màn hình đăng nhập, bấm Khám phá trường (Chế độ Khách)."),
            ("Bản đồ", "Có thể xem bản đồ GPS hoặc mô hình 3D khuôn viên trường."),
            ("Giới thiệu trường", "Xem thông tin trường, gallery ảnh, chiến dịch Dam San Green và bản tin."),
            ("Trợ lý AI", "Bấm bong bóng trợ lý để hỏi nhanh về trường, nội trú và phong trào xanh."),
        ],
    )

    add_heading(doc, "6. Đăng nhập và sử dụng Admin", 1)
    add_note_box(
        doc,
        "Cách mở cửa admin",
        "Ở màn hình đăng nhập, chạm vào logo ứng dụng 7 lần liên tiếp trong khoảng 2,5 giây. Hộp đăng nhập admin sẽ hiện ra. Nhập email/mật khẩu admin đã tạo trong Firebase và tài khoản đó phải có role = admin tại /Users/{uid}/role.",
    )
    add_step_table(
        doc,
        "Luồng admin",
        [
            ("Mở cửa admin", "Chạm logo 7 lần"),
            ("Đăng nhập", "Email admin + mật khẩu admin"),
            ("Duyệt báo cáo", "Chọn chờ duyệt/đã duyệt/từ chối"),
            ("Cài đặt", "Cập nhật bản tin, ảnh, logo, reset dữ liệu test"),
        ],
    )
    add_two_col_table(
        doc,
        "Chức năng trong Admin Panel",
        [
            ("Duyệt rác", "Xem báo cáo học sinh gửi, ảnh minh chứng, trạng thái AI; bấm Duyệt hoặc Từ chối."),
            ("Bảng xếp hạng", "Mở bảng vàng để xem lớp có tổng điểm cao nhất."),
            ("Bản tin trường", "Vào Cài đặt/Bản tin, nhập tiêu đề, nội dung, chọn ảnh rồi bấm Thêm bản tin và Lưu cài đặt."),
            ("Ảnh trường", "Upload ảnh giới thiệu lên Cloudinary thông qua nút chọn ảnh trong phần cài đặt."),
            ("Reset dữ liệu test", "Dùng khi kết thúc giai đoạn thử nghiệm: xóa báo cáo/điểm test, giữ tài khoản và branding."),
        ],
    )

    doc.add_page_break()
    add_heading(doc, "7. Xử lý lỗi thường gặp", 1)
    issues = [
        ("Không đăng nhập được", "Kiểm tra đúng lớp, đúng mật khẩu, thiết bị có Internet và tài khoản đã tồn tại trong Firebase Auth."),
        ("Báo cáo hiện lớp Unknown", "Cài lại APK mới nhất. Nếu vẫn lỗi, kiểm tra email tài khoản lớp có đúng dạng 11a1@damsan.edu.vn không."),
        ("Không upload được ảnh", "Kiểm tra mạng, quyền đọc ảnh/camera, cấu hình Cloudinary và dung lượng ảnh."),
        ("Không lưu được cài đặt admin", "Kiểm tra Realtime Database Rules có cho /Settings ghi khi /Users/{uid}/role = admin."),
        ("Không lấy được GPS", "Bật vị trí trên máy, cấp quyền vị trí cho app; nếu chạy máy ảo hãy đặt location trong emulator."),
    ]
    add_two_col_table(doc, "Bảng lỗi nhanh", issues)

    doc.add_page_break()
    add_heading(doc, "8. Bảng tài khoản đăng nhập", 1)
    add_note_box(
        doc,
        "Bảo mật khi in tài khoản",
        "Chỉ điền mật khẩu thật vào bản in/USB nộp cho giáo viên hoặc ban giám khảo. Không gửi bảng mật khẩu vào nhóm học sinh nếu chưa đổi mật khẩu riêng cho từng lớp.",
        fill="FFF1F0",
        accent=RED,
    )

    admin_rows = [
        ("Cửa admin", "Chạm logo ở màn hình đăng nhập 7 lần liên tiếp."),
        ("Email admin", "........................................................ hoặc admin@damsan.edu.vn nếu đã tạo theo checklist."),
        ("Mật khẩu admin", "........................................................"),
        ("Điều kiện quyền", "Realtime Database: /Users/{uid}/role = admin."),
    ]
    add_two_col_table(doc, "Tài khoản admin", admin_rows)

    for grade in (10, 11, 12):
        add_account_table(doc, grade)

    add_heading(doc, "9. Checklist trước khi demo/nộp", 1)
    checklist = [
        ("Cài APK mới nhất", "Đảm bảo dùng bản đã sửa lỗi class Unknown và lưu cài đặt admin."),
        ("Kiểm tra 1 tài khoản lớp", "Đăng nhập 11A1, chụp thử ảnh rác, gửi báo cáo, xem lớp hiển thị đúng."),
        ("Kiểm tra admin", "Vào cửa admin, duyệt báo cáo, mở bảng xếp hạng."),
        ("Kiểm tra bản tin", "Thêm 1 bản tin có ảnh, lưu cài đặt, xem lại trong trang giới thiệu trường."),
        ("Chuẩn bị mạng", "Khi demo cần Internet ổn định để Firebase, Cloudinary, bản đồ và AI hoạt động."),
    ]
    add_two_col_table(doc, "Danh sách kiểm tra", checklist)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Dam San Green - Hướng dẫn sử dụng và tài khoản đăng nhập")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
