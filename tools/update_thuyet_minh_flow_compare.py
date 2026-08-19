import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D6E8DD", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_run(run, size=10.5, bold=False, color="1B263B", italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if style:
        result.style = style
    if text:
        result.add_run(text)
    return result


def paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if style:
        result.style = style
    if text:
        result.add_run(text)
    return result


def table_after(document, paragraph, rows, cols):
    table = document.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(table._tbl)
    return table


def table_after_table(document, anchor_table, rows, cols):
    table = document.add_table(rows=rows, cols=cols)
    anchor_table._tbl.addnext(table._tbl)
    return table


def add_flowchart(document, anchor):
    caption = paragraph_after(anchor, "Sơ đồ hoạt động tổng quát của Dam San Green", style=anchor.style)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    for run in caption.runs:
        format_run(run, size=10.5, bold=True, color="146B4A")

    table = table_after(document, caption, 1, 7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 9000)

    labels = [
        "Học sinh chụp ảnh rác",
        "→",
        "Lưu trữ Firebase / Cloudinary",
        "→",
        "AI phân tích & dự đoán",
        "→",
        "Admin duyệt & cộng điểm",
    ]
    widths = [1800, 360, 2100, 360, 1800, 360, 2220]
    colors = ["E9FFF4", "F0FBFF", "FFFBE6", "EAF7EF"]

    for index, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[index])
        set_cell_border(cell, color="CBEBDD", size="8")
        set_cell_margins(cell, top=150, bottom=150, start=110, end=110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "FFFFFF" if labels[index] == "→" else colors[index // 2])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(labels[index])
        format_run(
            run,
            size=10.2 if labels[index] != "→" else 15,
            bold=True,
            color="146B4A" if labels[index] != "→" else "00A86B",
        )

    note = Paragraph(OxmlElement("w:p"), anchor._parent)
    table._tbl.addnext(note._p)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run(
        "Luồng này cho thấy công nghệ được dùng để biến hoạt động vệ sinh hằng ngày "
        "thành dữ liệu có minh chứng và điểm thi đua realtime."
    )
    format_run(run, size=9, italic=True, color="5F6C7B")


def add_compare_table(document, section_6_anchor, section_5_anchor):
    heading = paragraph_before(
        section_6_anchor,
        "Bảng so sánh trước và sau khi áp dụng Dam San Green",
        style=section_5_anchor.style,
    )
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(6)
    for run in heading.runs:
        format_run(run, size=11, bold=True, color="146B4A")

    table = table_after(document, heading, 5, 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 9000)

    headers = ["Trước đây (Truyền thống)", "Sau khi dùng Dam San Green"]
    rows = [
        (
            "Ghi nhận dọn rác bằng sổ cờ đỏ/phiếu giấy; học sinh dễ quên nhiệm vụ.",
            "Số hóa trên app Android; nhiệm vụ, ảnh minh chứng, thời gian và tài khoản lớp được lưu rõ ràng.",
        ),
        (
            "Mất thời gian tổng hợp điểm, dễ thiếu sót khi nhiều lớp cùng tham gia.",
            "Điểm thi đua cập nhật tự động sau khi duyệt; bảng xếp hạng realtime tạo động lực.",
        ),
        (
            "Khó xác minh rác đã dọn ở đâu, khi nào, có đúng khu vực phụ trách hay không.",
            "Minh chứng bằng ảnh, GPS và lịch sử báo cáo; AI hỗ trợ phân loại nhanh để admin kiểm tra.",
        ),
    ]

    for column, header in enumerate(headers):
        cell = table.cell(0, column)
        set_cell_width(cell, 4500)
        set_cell_shading(cell, "146B4A" if column == 0 else "00A86B")
        set_cell_border(cell, color="CBEBDD", size="8")
        set_cell_margins(cell, top=70, bottom=70, start=120, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        format_run(run, size=9.4, bold=True, color="FFFFFF")

    for row_index, row in enumerate(rows, start=1):
        for column, text in enumerate(row):
            cell = table.cell(row_index, column)
            set_cell_width(cell, 4500)
            set_cell_shading(cell, "FFFFFF" if row_index % 2 else "F6FBF8")
            set_cell_border(cell, color="CBEBDD", size="8")
            set_cell_margins(cell, top=45, bottom=45, start=100, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            format_run(run, size=8.8, color="1B263B")


def main():
    path = Path(sys.argv[1])
    backup = path.with_name(path.stem + ".backup_truoc_bo_sung_so_do_bang_so_sanh.docx")
    if not backup.exists():
        shutil.copy2(path, backup)

    document = Document(str(path))

    section_43 = None
    section_5 = None
    section_6 = None
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text.startswith("4.3."):
            section_43 = paragraph
        elif text.startswith("5.") and "Tính mới" in text:
            section_5 = paragraph
        elif text.startswith("6.") and "Kết quả" in text:
            section_6 = paragraph

    if not section_43 or not section_5 or not section_6:
        raise RuntimeError(
            f"Không tìm thấy đủ anchor: 4.3={bool(section_43)}, 5={bool(section_5)}, 6={bool(section_6)}"
        )

    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if "Sơ đồ hoạt động tổng quát của Dam San Green" not in all_text:
        add_flowchart(document, section_43)
    if "Bảng so sánh trước và sau khi áp dụng Dam San Green" not in all_text:
        add_compare_table(document, section_6, section_5)

    document.save(str(path))
    print(f"saved={path}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
